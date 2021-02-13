'''
    bwxref.py
    
    2/2/21  Ver 0.1
    2/4/21  Multiple Korean Bible DB
            Added BibleWorks Export 
    2/11/21 WTT Mapping
    
    2/13/21 Package not found at ... library.zip/docx/templates/default.docx
            1) pip install -U setuptools  (not working)
            2) "includes": ['lxml.etree', 'lxml._elementpath' ... (not working)
            3) copy default.docx and default-settings.xml to XRef folder
               Document('default.docx') (working!!)
'''
import re
import os
import lxml
import sqlite3 as db
import clipboard
import HTML 
import bwxreflib
import bwxrefwtt
import bwxrefcom
import wttmap as wmap
 
write_format = ['Txt', 'Docx', 'Html']

# Isa 66:5 dddd
# Eze 16:52-56 ddd
# Gen. 1:2,3,4f,5ff ddd

#_nkr_db_table_name = "Bible"
_xref_kbible_listfile = "kbible_list.txt"
_xref_ebible_listfile = "ebible_list.txt"

# (.*)(?<=\s)(\d+):(\d+[-\d]*)(,\d*f*)*
#_find_bwverse = re.compile('(.*)(?<=\s)(\d*):(\d*[-\d*]*)([,\d*]*)')
#_find_bwverse = re.compile('(.*)(?<=\s)(\d*):(\d*[-\d*]*)')
#_find_bwverse = re.compile('(.*)(?<=\s)(\d*):(\d*[-\d*]*)([,\s*\d*a-z]*)(?<=\s)')
_find_bwverse = re.compile('(.*)(?<=\s)(\d*):(\d*[-\d*]*)([,\da-z]*)')
    
class xref_elem:
    def __init__(self, book=None, chap=None, v1=0, v2=0):
        self.book = book
        self.chap = chap
        self.v1   = v1
        self.v2   = v2
        #self.vlist= []

xref_list = []
kbible_list = []
ebible_list = []
hgbible_list = ["Hebrew", "Greek"]
kbible_check_list = {}
ebible_check_list = {}
hgbible_check_list = {hgbible_list[0]:False, hgbible_list[1]:False}

def get_wtt_info(record):
    return record[0], record[1], "(WTT %d:%d)"%(record[0],record[1])
    
def get_wtt(map_table, key, chap, vers):
    wtt_map = bwxrefwtt.get_wtt_map_version(bwxrefwtt.get_wtt_map_key_wtt())
    map_key = '%02d:%d:%d'%(key, chap, vers)
    map_chap, map_vers, map_msg = chap, vers, ''
    
    if map_key in wtt_map:
        map_chap, map_vers, map_msg = get_wtt_info(wtt_map[map_key])
    elif map_key in map_table:
        map_chap, map_vers, map_msg = get_wtt_info(map_table[map_key])

    return map_chap, map_vers, map_msg
        
def write_kbible_list():
    
    fo = open(_xref_kbible_listfile, 'wt')
    fo.write('%d\n'%len(bwxreflib.korean_bible_name))
    
    for eb in bwxreflib.korean_bible_name:
        fo.write('%s\n'%eb)
    fo.close()
    
    
def read_kbible_list(): 
    global kbible_list, kbible_check_list
    
    try:
        fr = open(_xref_kbible_listfile, 'rt')
    except:
        write_kbible_list()
        fr = open(_xref_kbible_listfile, 'rt')
    
    try:
        nb = int(fr.readline())
    except Exception as e:
        print('... Error(read_kbible_list) => %s corrupted'% _xref_kbible_listfile)
        fr.close()
        write_kbible_list()
        fr = open(_xref_kbible_listfile, 'rt')
        nb = int(fr.readline())
    
    ebible_list = []
    ebible_check_list = {}
    
    for i in range(nb):
        kbn = fr.readline().strip()
        kbible_list.append(kbn)
        kbible_check_list[kbn] = True if kbn == "개역개정" else False
    fr.close()
                
def write_ebible_list():
    
    fo = open(_xref_ebible_listfile, 'wt')
    fo.write('%d\n'%len(bwxreflib.english_bible_name))
    
    for kb in bwxreflib.english_bible_name:
        fo.write('%s\n'%kb)
    fo.close()
    
def read_ebible_list(): 
    global ebible_list, ebible_check_list
    
    try:
        fr = open(_xref_ebible_listfile, 'rt')
    except:
        write_ebible_list()
        fr = open(_xref_ebible_listfile, 'rt')
    
    try:
        nb = int(fr.readline())
    except Exception as e:
        print('... Error(read_ebible_list) => %s corrupted'% _xref_ebible_listfile)
        fr.close()
        write_ebible_list()
        fr = open(_xref_ebible_listfile, 'rt')
        nb = int(fr.readline())
    ebible_list = []
    ebible_check_list = {}
         
    for i in range(nb):
        ebn = fr.readline().strip()
        ebible_list.append(ebn)
        ebible_check_list[ebn] = False
    fr.close()
    
def get_kbible_list():
    return kbible_list

def get_ebible_list():
    return ebible_list
  
def get_hgbible_list():
    return hgbible_list
    
def get_kbible_check_list():
    return kbible_check_list

def get_ebible_check_list():
    return ebible_check_list
 
def get_hgbible_check_list():
    return hgbible_check_list
     
def xref_to_txt(path, file, xref, wtt_map, map_table, db_list, msg):
    
    file = '%s.txt'%file
    msg.appendPlainText('... XRef to Txt\n... Open %s'%file)
    try:
        fo = open(os.path.join(path, file), 'wt')
    except Exception as e:
        e_str = str(e)
        msg.appendPlainText('... Error(xref_to_txt) ==> %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return
    
    for x in xref:
        e = bwxreflib.table["%02d"%x.book]
        for key in db_list.keys():
            db_file = db_list[key]
            try:
                msg.appendPlainText('... Open DB %s'%db_file)
                db_con = db.connect(db_file)
                db_cur = db_con.cursor()
                table  = db_cur.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()
                db_tbl = table[0][0]
                fo.write('====== %s ======\n'%key)
            except Exception as e:
                e_str = str(e)
                msg.appendPlainText('... DB Error => %s'%e_str)
                bwxrefcom.message_box(bwxrefcom.message_error, e_str)
                return
            
            if x.v2 is 0:
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, x.v1)
                else:
                    chap, vers, m_msg = x.chap, x.v1, ''

                sql = 'SELECT btext from %s where book=%d '\
                      'and chapter=%d and verse=%d'%\
                      (db_tbl, x.book, chap, vers)
                      #(db_tbl, x.book, x.chap, x.v1)
                      
                fo.write('%s %d:%d\n'%(e[2], x.chap, x.v1))
                
                for vtext in db_cur.execute(sql):
                    fo.write('%d. %s%s\n'%(x.v1, vtext[0], m_msg))
            else:
                v_list = [vv for vv in range(x.v1, x.v2+1)]
                
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    for vv in v_list:
                        chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, vv)
                        sql = 'SELECT btext from %s where book=%d '\
                              'and chapter=%d and verse=%d'%\
                              (db_tbl, x.book, chap, vv)
                        v_cur = db_cur.execute(sql)
                        fo.write('%d. %s%s'%(vv, v_cur.fetchone()[0], m_msg))
                else:
                    sql = 'SELECT btext from %s where book=%d '\
                        'and chapter=%d and verse between %d and %d'%\
                        (db_tbl, x.book, x.chap, x.v1, x.v2)
                      
                    fo.write('%s %d:%d-%d\n'%(e[2], x.chap, x.v1, x.v2))
                    v_cur = db_cur.execute(sql)
                    for vi in v_list:
                        fo.write('%d. %s\n'%(x.v1+i, v_cur.fetchone()[0]))
        fo.write('\n')
        db_con.close()
        msg.appendPlainText('... Close DB')
    fo.write('\n')
    
    msg.appendPlainText('... Success')
    bwxrefcom.message_box(bwxrefcom.message_normal, 'Success')

def xref_to_docx(path, file, xref, wtt_map, map_table, db_list, msg):
    from docx import Document
    from docx.shared import Inches

    file = '%s.docx'%file
    msg.appendPlainText('... XRef to Docx\n... Open: %s'%file)

    ndb = len(db_list)
    
    try:
        document = Document('default.docx')
    except Exception as e:
        e_str = "... Error(xref_to_docx): Can't create Word Document.\n%s"%str(e)
        msg.appendPlainText(e_str)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return
        
    rows = len(xref)+1
    cols = ndb+1
    msg.appendPlainText('... Table(row,col): %d x %d'%(rows,cols))
    table = document.add_table(rows=rows, cols=cols)
    
    table.rows[0].cells[0].text = "성경구절"
    
    for icol, (key, db_file) in enumerate(db_list.items()):
        icol += 1
        table.rows[0].cells[icol].text = key
        
        try:
            msg.appendPlainText('... Open DB %s'%db_file)
            db_con = db.connect(db_file)
            db_cur = db_con.cursor()
            db_tbls= db_cur.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()
            db_tbl = db_tbls[0][0]
        except Exception as e:
            e_str = str(e)
            #print('Error =>', str(e))
            msg.appendPlainText('... DB Error => %s'%e_str)
            bwxrefcom.message_box(bwxrefcom.message_error, e_str)
            return
        
            
        for irow, x in enumerate(xref):
            irow += 1
            #table.rows[irow].cells[0]._tc.tcPr.tcW.type = 'auto'
            #table.rows[irow].cells[0]._tc.tcPr.tcW.w = 0
        
            e = bwxreflib.table["%02d"%x.book]
            #table.rows[0].cells[icol].text = '%s %d:%d'%(e[2], x.chap, x.v1) if x.v2 == 0 else\
            #                                 '%s %d:%d-d%'%(e[2], x.chap, x.v1, x.v2)
            if x.v2 == 0:
                table.rows[irow].cells[0].text = '%s %d:%d'%(e[2], x.chap, x.v1)
            else:
                table.rows[irow].cells[0].text = '%s %d:%d-%d'%(e[2], x.chap, x.v1, x.v2)          
            
            if x.v2 is 0:
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, x.v1)
                else:
                    chap, vers, m_msg = x.chap, x.v1, ''
                sql = 'SELECT btext from %s where book=%d '\
                      'and chapter=%d and verse=%d'%\
                      (db_tbl, x.book, chap, vers)
                          #(db_tbl, x.book, x.chap, x.v1)
                v_cur = db_cur.execute(sql)
                table.rows[irow].cells[icol].text = '%d. %s%s'%(x.v1, v_cur.fetchone()[0], m_msg)
            else:
                table.rows[irow].cells[0].text = '%s %d:%d-%d'%(e[2], x.chap, x.v1, x.v2)
                v_list = [vv for vv in range(x.v1, x.v2+1)]
                vl = []
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    for vv in v_list:
                        chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, vv)
                        sql = 'SELECT btext from %s where book=%d '\
                              'and chapter=%d and verse=%d'%\
                              (db_tbl, x.book, chap, vv)
                        v_cur = db_cur.execute(sql)
                        vl.append('<p>%d. %s%s</p>'%(vv, v_cur.fetchone()[0], m_msg))
                else:
                    sql = 'SELECT btext from %s where book=%d '\
                          'and chapter=%d and verse between %d and %d'%\
                          (db_tbl, x.book, x.chap, x.v1, x.v2)
                    v_cur = db_cur.execute(sql)
                    for vi in v_list:
                        vl.append('%d. %s'%(vi, v_cur.fetchone()[0]))
                table.rows[irow].cells[icol].text = '\n'.join(vl)
    
    try:
        document.save(os.path.join(path,file))
    except Exception as e:
        e_str = str(e)
        msg.appendPlainText('... Error(xref_to_docx) ==> %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return
        
    msg.appendPlainText('... success')
    bwxrefcom.message_box(bwxrefcom.message_normal, 'Success')

# map_table is a tuple (key, map_dictionary)

def xref_to_html(path, file, xref, wtt_map, map_table, db_list, msg):
    file = '%s.html'%file
    msg.appendPlainText('... XRef to Html\n... Open %s'%file)
    try:
        fo = open(os.path.join(path, file), mode='wt', encoding='utf8')
    except Exception as e:
        e_str = str(e)
        msg.appendPlainText('... Error ==> %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return
    
    head = [key for key in db_list.keys()]
    head.insert(0, ' ')
    html_table = HTML.Table(header_row=head)
    
    for x in xref:
        e = bwxreflib.table["%02d"%x.book]
        if x.v2 is 0:
            html_row = ['%s %d:%d'%(e[2], x.chap, x.v1)]
        else:
            html_row = ['%s %d:%d-%d\n'%(e[2], x.chap, x.v1, x.v2)]

        for col, (key, db_file) in enumerate(db_list.items()):
            db_file = db_list[key]
            try:
                msg.appendPlainText('... Open DB %s'%db_file)
                db_con = db.connect(db_file)
                db_cur = db_con.cursor()
                table  = db_cur.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()
                db_tbl = table[0][0]
                #msg.appendPlainText(db_tbl)
            except Exception as e:
                e_str = str(e)
                msg.appendPlainText('... DB Error => %s'%e_str)
                bwxrefcom.message_box(bwxrefcom.message_error, e_str)
                return

            if x.v2 is 0:
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, x.v1)
                else:
                    chap, vers, m_msg = x.chap, x.v1, ''
                sql = 'SELECT btext from %s where book=%d '\
                      'and chapter=%d and verse=%d'%\
                      (db_tbl, x.book, chap, vers)
                      #(db_tbl, x.book, x.chap, x.v1)
                      
                v_cur = db_cur.execute(sql)
                html_row.append('%d. %s%s'%(x.v1, v_cur.fetchone()[0], m_msg))
                #for vtext in db_cur.execute(sql):
                #    html_row.append('%d. %s%s'%(x.v1, vtext[0], m_msg))
            else:
                v_list = [vv for vv in range(x.v1, x.v2+1)]
                if wtt_map and x.book <= 39 and map_table[0] != bwxrefwtt.get_wtt_map_key_nau(): 
                    vl = []
                    for vv in v_list:
                        chap, vers, m_msg = get_wtt(map_table[1], x.book, x.chap, vv)
                        sql = 'SELECT btext from %s where book=%d '\
                              'and chapter=%d and verse=%d'%\
                              (db_tbl, x.book, chap, vv)
                        v_cur = db_cur.execute(sql)
                        vl.append('<p>%d. %s%s</p>'%(vv, v_cur.fetchone()[0], m_msg))
                        #for vtext in db_cur.execute(sql):
                        #    vl.append('<p>%d. %s%s</p>'%(vv, vtext[0], m_msg))
                else:
                    sql = 'SELECT btext from %s where book=%d '\
                        'and chapter=%d and verse between %d and %d'%\
                        (db_tbl, x.book, x.chap, x.v1, x.v2)
                    v_cur = db_cur.execute(sql)
                    vl = []
                    for vi in v_list:
                        vl.append('<p>%d. %s</p>'%(vi, v_cur.fetchone()[0]))
                html_row.append(''.join(vl))
            db_con.close()
            msg.appendPlainText('... Close DB')
        html_table.rows.append(html_row)

    fo.write(str(html_table))
    msg.appendPlainText('... Success')
    bwxrefcom.message_box(bwxrefcom.message_normal, 'Success')
    
write_xref = { 
    write_format[0] : xref_to_txt,  
    write_format[1] : xref_to_docx,  
    write_format[2] : xref_to_html,  
}

def get_write_format_txt():
    return write_format[0]
    
def get_write_format_docx():
    return write_format[1]

def get_write_format_html():
    return write_format[2]
    
def xref_to_kor(path, 
                file, 
                type, 
                fmt=write_format[0], 
                wtt_map = False,
                map_table=None, 
                db_list=None, 
                msg=None):
                
    global xref_list
   
    msg.appendPlainText('... xref_to_kor\n... Get clipboard data')
    try:
       verselist = clipboard.paste().split('\n')
    except Exception as e:
        #print('Error =>', str(e))
        e_str = str(e)
        msg.appendPlainText('... Error(xref_to_txt) => %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return
    
    if not verselist: 
        e_str = 'Error => No clipboard data!'
        msg.appendPlainText('... %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_warning, e_str)
        return
    
    xref_list = []
    non_canon_book = False
    
    for v in verselist:
        match = _find_bwverse.search(v)
        if match:
            book = match.group(1).strip()
            num  = bwxreflib.get_book_num(book, type)
            if num < 0:
                msg.appendPlainText('... Non-canon book: %s'%book)
                non_canon_book = True
                continue
                
            ngroup= len(match.groups())
            book  = int(num)
            chap  = int(match.group(2))
            verse = match.group(3)
            extra = match.group(4)
            
            if verse.find('-') > -1:
                vv = verse.split('-')
                v1 = int(vv[0])
                v2 = int(vv[1])
                xref = xref_elem(book, chap, v1, v2)
            elif extra and extra.find(',') > -1:
                # example: Matt. 24:3,27f,37,39ff
                xref_list.append(xref_elem(book, chap, int(verse)))
                vv = extra.split(',')
                vv.pop(0)
                for v1 in vv:
                    v2 = re.sub('[a-z]*', '', v1.strip())
                    xref_list.append(xref_elem(book, chap, int(v2)))
                continue
            else:
                xref = xref_elem(book, chap, int(verse))
                
            xref_list.append(xref)

    if len(xref_list) is 0:
        e_str = 'Error => No verse exist! Check Clipboard!'
        msg.appendPlainText('... %s'%e_str)
        bwxrefcom.message_box(bwxrefcom.message_warning, e_str)
        return
        
    write_xref[fmt](path, file, xref_list, wtt_map, map_table, db_list, msg)
    
    if non_canon_book: 
        bwxrefcom.message_box(bwxrefcom.message_normal, "Check the message")
    msg.appendPlainText('... Success')
    
#xref_to_kor("xref.docx", write_format[1])
#xref_to_kor("xref.txt", write_format[0])
